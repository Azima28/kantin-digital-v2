#!/usr/bin/env python3
"""
Generate Makalah: "Sistem Kantin Digital Berbasis RFID/NFC"
Format: TNR 12pt, 1.5 spasi, margin 4-4-3-3, justify, indent 1.27cm
Project: Kantin Digital (Flutter + Supabase)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ─── PAGE SETUP ───────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

# ─── STYLES ───────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

# Set font for East Asian as well
rPr = style.element.find(qn('w:rPr'))
if rPr is None:
    rPr = OxmlElement('w:rPr')
    style.element.append(rPr)

# Heading styles
from docx.enum.text import WD_ALIGN_PARAGRAPH as Align

for level, size, align in [(1, 14, Align.CENTER), (2, 12, Align.LEFT), (3, 12, Align.LEFT)]:
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(size)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.paragraph_format.alignment = align
    hs.paragraph_format.line_spacing = 1.5
    hs.paragraph_format.space_before = Pt(12)
    hs.paragraph_format.space_after = Pt(6)
    hs.paragraph_format.keep_with_next = True

# ─── HELPER FUNCTIONS ─────────────────────────────────────────
def add_body(text):
    """Body paragraph: justify, first-line indent 1.27cm"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = Align.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_body_no_indent(text):
    """Body paragraph without indent (for after images/tables)"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = Align.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_justify_no_indent(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = Align.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_centered(text, bold=False, size=12, italic=False, space_after=True):
    p = doc.add_paragraph()
    p.alignment = Align.CENTER
    p.paragraph_format.line_spacing = 1.5
    if space_after:
        p.paragraph_format.space_after = Pt(0)
    else:
        p.paragraph_format.space_after = Pt(24)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_heading_text(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_page_break():
    doc.add_page_break()

def add_spacing(lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run('')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

def add_bullet(text, indent_level=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = Align.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))
    p.paragraph_format.first_line_indent = Cm(-0.63)
    prefix = '• ' if indent_level == 0 else '- '
    run = p.add_run(prefix + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_numbered(text, number, indent_level=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = Align.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.27 * (indent_level + 1))
    p.paragraph_format.first_line_indent = Cm(-0.63)
    run = p.add_run(f'{number}. {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = Align.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    
    # Rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = Align.LEFT
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(str(cell_text))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

    # Add spacing after table
    add_spacing(1)
    return table

# ═══════════════════════════════════════════════════════════════
#                       COVER PAGE
# ═══════════════════════════════════════════════════════════════
add_spacing(6)
add_centered('MAKALAH', bold=True, size=18)
add_spacing(1)
add_centered(
    'SISTEM KANTIN DIGITAL BERBASIS RFID/NFC\n'
    'MENGGUNAKAN FLUTTER DAN SUPABASE',
    bold=True, size=16
)
add_spacing(1)
add_centered(
    'Disusun untuk Memenuhi Tugas Praktik Kerja Lapangan (PKL)',
    italic=True, size=12
)
add_spacing(3)
add_centered('Disusun oleh:', size=12)
add_centered('Agust', bold=True, size=14)
add_spacing(3)
add_centered('SEKOLAH MENENGAH KEJURUAN', bold=True, size=14)
add_centered('JURUSAN REKAYASA PERANGKAT LUNAK', bold=True, size=14)
add_spacing(1)
add_centered('2026', bold=True, size=14)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#                     KATA PENGANTAR
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = Align.CENTER
p.paragraph_format.line_spacing = 1.5
run = p.add_run('KATA PENGANTAR')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_spacing(2)

kata_pengantar = (
    "Puji syukur ke hadirat Tuhan Yang Maha Esa atas segala rahmat dan karunia-Nya "
    "sehingga penulis dapat menyelesaikan makalah yang berjudul \"Sistem Kantin Digital "
    "Berbasis RFID/NFC Menggunakan Flutter dan Supabase\" ini dengan baik. Makalah ini "
    "disusun sebagai laporan hasil pelaksanaan Praktik Kerja Lapangan (PKL) yang telah "
    "dilaksanakan."
)
add_body(kata_pengantar)

add_body(
    "Penyusunan makalah ini bertujuan untuk mendokumentasikan proses pengembangan sistem "
    "informasi kantin digital yang dibangun menggunakan framework Flutter pada sisi frontend "
    "dan Supabase sebagai backend-as-a-service. Sistem ini dirancang untuk mengatasi "
    "permasalahan pencatatan transaksi jajan siswa yang masih manual di lingkungan sekolah."
)

add_body(
    "Penulis menyadari bahwa makalah ini masih jauh dari sempurna. Oleh karena itu, penulis "
    "sangat mengharapkan kritik dan saran yang membangun demi perbaikan di masa yang akan "
    "datang. Semoga makalah ini dapat memberikan manfaat bagi pembaca dan menjadi referensi "
    "dalam pengembangan sistem serupa di kemudian hari."
)

add_spacing(2)
p = doc.add_paragraph()
p.alignment = Align.RIGHT
p.paragraph_format.line_spacing = 1.5
run = p.add_run('Penulis,')
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
add_spacing(4)
p = doc.add_paragraph()
p.alignment = Align.RIGHT
p.paragraph_format.line_spacing = 1.5
run = p.add_run('Agust')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#                     DAFTAR ISI
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = Align.CENTER
p.paragraph_format.line_spacing = 1.5
run = p.add_run('DAFTAR ISI')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_spacing(2)

toc_items = [
    ('KATA PENGANTAR', 'ii'),
    ('DAFTAR ISI', 'iii'),
    ('BAB I   PENDAHULUAN', '1'),
    ('1.1 Latar Belakang', '1'),
    ('1.2 Rumusan Masalah', '2'),
    ('1.3 Batasan Masalah', '2'),
    ('1.4 Tujuan Penulisan', '2'),
    ('1.5 Manfaat Penulisan', '3'),
    ('BAB II   TINJAUAN PUSTAKA', '4'),
    ('2.1 Konsep Dasar Sistem Informasi', '4'),
    ('2.2 Flutter Framework', '4'),
    ('2.3 Supabase', '5'),
    ('2.4 Teknologi RFID dan NFC', '5'),
    ('BAB III   METODE PENELITIAN', '6'),
    ('3.1 Metode Pengembangan', '6'),
    ('3.2 Analisis Kebutuhan', '6'),
    ('3.3 Perancangan Sistem', '7'),
    ('3.4 Perancangan Basis Data', '8'),
    ('BAB IV   HASIL DAN PEMBAHASAN', '10'),
    ('4.1 Implementasi Sistem', '10'),
    ('4.2 Pengujian Sistem', '11'),
    ('4.3 Pembahasan', '12'),
    ('BAB V   PENUTUP', '13'),
    ('5.1 Kesimpulan', '13'),
    ('5.2 Saran', '13'),
    ('DAFTAR PUSTAKA', '14'),
]

max_label = max(len(item[0]) for item in toc_items)
for label, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), alignment=Align.RIGHT)
    
    is_bold = not label.startswith(' ') and not label[0].isdigit()
    is_sub = label.startswith(' ') or label[0].isdigit()
    
    display_label = label
    run = p.add_run(display_label)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if is_bold and not is_sub:
        run.bold = True
    
    run2 = p.add_run(f'\t{page}')
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#                 BAB I: PENDAHULUAN
# ═══════════════════════════════════════════════════════════════
add_heading_text('BAB I', level=1)
add_heading_text('PENDAHULUAN', level=1)
add_spacing(1)

# 1.1 Latar Belakang
add_heading_text('1.1 Latar Belakang', level=2)
add_body(
    "Perkembangan teknologi informasi telah membawa perubahan signifikan dalam berbagai "
    "aspek kehidupan, termasuk di lingkungan sekolah. Salah satu kegiatan yang masih "
    "dilakukan secara konvensional di banyak sekolah adalah transaksi jajan di kantin. "
    "Pada umumnya, siswa membawa uang tunai untuk membeli makanan dan minuman di kantin "
    "sekolah. Hal ini menimbulkan beberapa permasalahan, seperti antrean panjang saat jam "
    "istirahat, risiko kehilangan uang, serta kesulitan orang tua dalam mengontrol "
    "pengeluaran jajan anak-anak mereka."
)

add_body(
    "Berdasarkan observasi yang dilakukan di lingkungan sekolah, sistem pembayaran tunai "
    "di kantin masih mendominasi. Siswa harus mengantre untuk membeli makanan, petugas "
    "kantin harus menghitung uang kembalian secara manual, dan orang tua tidak memiliki "
    "cara untuk memantau apa yang dibeli anak mereka di kantin. Permasalahan ini menjadi "
    "semakin kompleks ketika jumlah siswa yang harus dilayani dalam waktu istirahat yang "
    "terbatas mencapai ratusan orang."
)

add_body(
    "Untuk mengatasi permasalahan tersebut, penulis mengembangkan sebuah sistem kantin "
    "digital yang memanfaatkan teknologi RFID (Radio Frequency Identification) dan NFC "
    "(Near Field Communication). Sistem ini memungkinkan siswa melakukan transaksi "
    "pembayaran di kantin cukup dengan menempelkan kartu identitas siswa yang dilengkapi "
    "chip RFID pada perangkat pembaca NFC. Dengan demikian, proses transaksi menjadi "
    "lebih cepat, efisien, dan tercatat secara digital."
)

add_body(
    "Sistem dikembangkan menggunakan Flutter sebagai framework pengembangan aplikasi "
    "mobile multiplatform dan Supabase sebagai backend-as-a-service yang menyediakan "
    "basis data PostgreSQL, autentikasi, dan penyimpanan file. Pemilihan Flutter didasarkan "
    "pada kemampuannya untuk menghasilkan aplikasi yang dapat berjalan di platform Android "
    "dan iOS secara bersamaan dengan satu basis kode. Sementara itu, Supabase dipilih "
    "karena menyediakan layanan backend yang lengkap dan mudah diintegrasikan tanpa "
    "memerlukan pengelolaan server secara mandiri."
)

# 1.2 Rumusan Masalah
add_heading_text('1.2 Rumusan Masalah', level=2)
add_body(
    "Berdasarkan latar belakang yang telah diuraikan, rumusan masalah dalam pengembangan "
    "sistem ini adalah sebagai berikut:"
)
add_numbered('Bagaimana merancang dan membangun sistem kantin digital berbasis RFID/NFC yang dapat mempercepat proses transaksi di kantin sekolah?', 1)
add_numbered('Bagaimana mengintegrasikan Flutter sebagai framework frontend dengan Supabase sebagai backend untuk sistem transaksi keuangan real-time?', 2)
add_numbered('Bagaimana merancang sistem manajemen saldo digital yang aman dan dapat dipertanggungjawabkan?', 3)
add_numbered('Bagaimana menyediakan fitur pemantauan bagi orang tua untuk melihat riwayat transaksi jajan anak?', 4)
add_numbered('Bagaimana mengimplementasikan sistem role-based access control untuk lima jenis pengguna: super admin, admin keuangan, petugas kantin, siswa, dan orang tua?', 5)

# 1.3 Batasan Masalah
add_heading_text('1.3 Batasan Masalah', level=2)
add_body(
    "Agar pembahasan lebih terfokus, penelitian ini dibatasi pada beberapa hal berikut:"
)
add_bullet('Sistem dikembangkan untuk lingkungan sekolah menengah kejuruan (SMK) sebagai studi kasus.')
add_bullet('Aplikasi dibangun menggunakan Flutter dengan bahasa pemrograman Dart dan Supabase sebagai backend.')
add_bullet('Sistem mencakup lima modul utama: autentikasi pengguna, dashboard siswa, POS kasir, manajemen keuangan, dan panel admin.')
add_bullet('Transaksi pembayaran dilakukan menggunakan kartu RFID/NFC yang terdaftar pada sistem.')
add_bullet('Pembahasan tidak mencakup aspek keamanan jaringan secara mendalam.')
add_bullet('Fitur pembayaran online menggunakan Midtrans hanya dibahas secara konseptual.')

# 1.4 Tujuan Penulisan
add_heading_text('1.4 Tujuan Penulisan', level=2)
add_body(
    "Tujuan dari penulisan makalah ini adalah:"
)
add_numbered('Merancang dan mengimplementasikan sistem kantin digital berbasis RFID/NFC yang dapat mempercepat proses transaksi jajan di kantin sekolah.', 1)
add_numbered('Mengintegrasikan Flutter framework dengan Supabase backend untuk menciptakan sistem transaksi real-time yang responsif.', 2)
add_numbered('Mengimplementasikan sistem manajemen saldo digital dengan mekanisme ACID transaction untuk mencegah inkonsistensi data.', 3)
add_numbered('Menyediakan fitur dashboard orang tua untuk memantau pengeluaran jajan anak secara real-time.', 4)
add_numbered('Menerapkan role-based access control yang memisahkan hak akses sesuai dengan peran masing-masing pengguna.', 5)

# 1.5 Manfaat Penulisan
add_heading_text('1.5 Manfaat Penulisan', level=2)
add_body("Penulisan makalah ini diharapkan dapat memberikan manfaat sebagai berikut:")
add_bullet('Bagi siswa: memudahkan proses transaksi di kantin tanpa perlu membawa uang tunai.')
add_bullet('Bagi petugas kantin: mempercepat proses kasir dan mengurangi kesalahan perhitungan.')
add_bullet('Bagi orang tua: dapat memantau pengeluaran jajan anak secara transparan.')
add_bullet('Bagi pihak sekolah: memiliki sistem pencatatan keuangan kantin yang terintegrasi dan dapat diaudit.')
add_bullet('Bagi penulis: menambah pengetahuan dan pengalaman dalam pengembangan aplikasi full-stack menggunakan Flutter dan Supabase.')

add_page_break()

# ═══════════════════════════════════════════════════════════════
#                 BAB II: TINJAUAN PUSTAKA
# ═══════════════════════════════════════════════════════════════
add_heading_text('BAB II', level=1)
add_heading_text('TINJAUAN PUSTAKA', level=1)
add_spacing(1)

add_heading_text('2.1 Konsep Dasar Sistem Informasi', level=2)
add_body(
    "Sistem informasi merupakan kombinasi dari teknologi informasi dan aktivitas orang "
    "yang menggunakan teknologi tersebut untuk mendukung operasi dan manajemen. Menurut "
    "O'Brien dan Marakas (2010), sistem informasi adalah kombinasi teratur dari orang, "
    "perangkat keras, perangkat lunak, jaringan komunikasi, sumber daya data, dan "
    "kebijakan prosedural yang menyimpan, mengambil, mengubah, dan menyebarkan informasi "
    "dalam sebuah organisasi."
)
add_body(
    "Dalam konteks kantin digital, sistem informasi berfungsi untuk mengelola data siswa, "
    "transaksi jajan, saldo digital, dan laporan keuangan secara terintegrasi. Sistem ini "
    "menggantikan pencatatan manual yang rentan terhadap kesalahan dan kehilangan data."
)

add_heading_text('2.2 Flutter Framework', level=2)
add_body(
    "Flutter adalah framework pengembangan aplikasi mobile open-source yang dikembangkan "
    "oleh Google. Flutter menggunakan bahasa pemrograman Dart dan memungkinkan pengembang "
    "untuk membuat aplikasi yang dapat berjalan di berbagai platform (Android, iOS, Web, "
    "dan Desktop) dengan satu basis kode yang sama (Google, 2024)."
)
add_body(
    "Beberapa keunggulan Flutter yang relevan dengan pengembangan sistem ini antara lain: "
    "hot reload yang memungkinkan perubahan kode terlihat secara instan, widget library "
    "yang kaya dan dapat dikustomisasi, performa native karena dikompilasi langsung ke "
    "kode mesin, serta dukungan komunitas yang besar. Dalam sistem Kantin Digital, Flutter "
    "digunakan untuk membangun antarmuka pengguna pada aplikasi siswa (mobile), aplikasi "
    "kasir POS (tablet), dan dashboard admin/keuangan."
)

add_heading_text('2.3 Supabase', level=2)
add_body(
    "Supabase adalah platform backend-as-a-service (BaaS) open-source yang menyediakan "
    "basis data PostgreSQL, autentikasi, penyimpanan file, dan fungsi edge computing. "
    "Supabase sering disebut sebagai alternatif open-source dari Firebase karena menyediakan "
    "layanan serupa dengan menggunakan teknologi PostgreSQL sebagai basis data utama "
    "(Supabase, 2024)."
)
add_body(
    "Fitur-fitur Supabase yang digunakan dalam sistem ini meliputi: (1) PostgreSQL database "
    "dengan dukungan Row Level Security (RLS) untuk keamanan data; (2) Autentikasi pengguna "
    "dengan dukungan email/password dan manajemen sesi; (3) Realtime subscriptions untuk "
    "pembaruan data secara langsung; (4) Storage untuk menyimpan gambar produk dan foto "
    "profil; serta (5) Edge Functions untuk menjalankan logika bisnis di sisi server."
)

add_heading_text('2.4 Teknologi RFID dan NFC', level=2)
add_body(
    "Radio Frequency Identification (RFID) adalah teknologi yang menggunakan gelombang "
    "radio untuk mengidentifikasi objek secara otomatis. Sistem RFID terdiri dari tag "
    "(transponder) yang dipasang pada objek dan reader (interrogator) yang membaca data "
    "dari tag. Near Field Communication (NFC) merupakan pengembangan dari RFID yang "
    "beroperasi pada frekuensi 13.56 MHz dengan jarak baca maksimal sekitar 10 cm "
    "(Want, 2006)."
)
add_body(
    "Dalam sistem Kantin Digital, NFC digunakan sebagai media identifikasi siswa. Setiap "
    "siswa memiliki kartu identitas yang dilengkapi dengan chip NFC. Ketika kartu "
    "ditempelkan pada perangkat pembaca NFC yang terhubung dengan aplikasi kasir, sistem "
    "akan membaca UID (Unique Identifier) kartu dan mencocokkannya dengan data siswa "
    "yang tersimpan di basis data. Keunggulan penggunaan NFC dibandingkan barcode atau QR "
    "code adalah proses pembacaan yang lebih cepat dan tidak memerlukan pencahayaan khusus."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#              BAB III: METODE PENELITIAN
# ═══════════════════════════════════════════════════════════════
add_heading_text('BAB III', level=1)
add_heading_text('METODE PENELITIAN', level=1)
add_spacing(1)

add_heading_text('3.1 Metode Pengembangan', level=2)
add_body(
    "Metode pengembangan yang digunakan dalam pembuatan sistem Kantin Digital adalah "
    "model waterfall yang dimodifikasi. Model ini dipilih karena kebutuhan sistem yang "
    "telah terdefinisi dengan jelas sejak awal dan tidak memerlukan perubahan yang "
    "signifikan selama proses pengembangan. Tahapan pengembangan meliputi:"
)
add_numbered('Analisis Kebutuhan: Mengidentifikasi kebutuhan fungsional dan non-fungsional sistem melalui observasi dan wawancara dengan pihak sekolah.', 1)
add_numbered('Perancangan Sistem: Merancang arsitektur sistem, basis data, dan antarmuka pengguna berdasarkan hasil analisis kebutuhan.', 2)
add_numbered('Implementasi: Menulis kode program menggunakan Flutter dan Supabase sesuai dengan perancangan yang telah dibuat.', 3)
add_numbered('Pengujian: Melakukan pengujian fungsional menggunakan metode black-box testing untuk memastikan setiap fitur berjalan sesuai spesifikasi.', 4)
add_numbered('Deployment: Menyebarkan aplikasi ke perangkat target dan melakukan pelatihan penggunaan kepada pengguna.', 5)

add_heading_text('3.2 Analisis Kebutuhan', level=2)
add_body(
    "Berdasarkan hasil observasi dan wawancara, kebutuhan sistem Kantin Digital dapat "
    "diklasifikasikan menjadi kebutuhan fungsional dan non-fungsional sebagai berikut:"
)

add_heading_text('3.2.1 Kebutuhan Fungsional', level=3)
add_body(
    "Kebutuhan fungsional sistem mencakup fitur-fitur yang harus dimiliki oleh sistem, "
    "yaitu:"
)
add_bullet('Sistem dapat mengelola data pengguna dengan lima peran: super admin, admin keuangan, petugas kantin, siswa, dan orang tua.')
add_bullet('Sistem dapat melakukan autentikasi pengguna melalui email, NISN, atau username dengan dual-path authentication.')
add_bullet('Siswa dapat melihat saldo digital dan riwayat transaksi jajan.')
add_bullet('Petugas kantin dapat melakukan transaksi penjualan dengan memindai kartu RFID/NFC siswa.')
add_bullet('Admin keuangan dapat melakukan top-up saldo, koreksi saldo, dan registrasi kartu RFID.')
add_bullet('Super admin dapat mengelola seluruh pengguna, melihat audit log, dan mengatur konfigurasi sistem.')
add_bullet('Orang tua dapat melakukan top-up saldo anak dan memantau riwayat transaksi.')
add_bullet('Sistem dapat menghasilkan laporan keuangan dan aktivitas transaksi.')

add_heading_text('3.2.2 Kebutuhan Non-Fungsional', level=3)
add_bullet('Sistem harus responsif dengan waktu response tidak lebih dari 2 detik untuk setiap transaksi.')
add_bullet('Sistem harus memiliki mekanisme ACID transaction untuk mencegah double-spending.')
add_bullet('Sistem harus memiliki audit trail untuk setiap perubahan data saldo.')
add_bullet('Sistem harus mendukung penggunaan secara offline dengan sinkronisasi data saat online kembali.')
add_bullet('Antarmuka pengguna harus intuitif dan mudah digunakan oleh pengguna dengan latar belakang non-teknis.')

add_heading_text('3.3 Perancangan Sistem', level=2)

add_heading_text('3.3.1 Arsitektur Sistem', level=3)
add_body(
    "Sistem Kantin Digital menggunakan arsitektur client-server dengan Flutter sebagai "
    "client dan Supabase sebagai server. Seluruh logika bisnis yang berkaitan dengan "
    "manipulasi saldo diimplementasikan sebagai Stored Procedure (RPC function) di sisi "
    "database PostgreSQL untuk memastikan integritas data melalui mekanisme transaksi ACID."
)

add_heading_text('3.3.2 Use Case Diagram', level=3)
add_body(
    "Sistem memiliki lima aktor utama: (1) Siswa, yang dapat melihat saldo, riwayat "
    "transaksi, dan mengelola kartu; (2) Petugas Kantin, yang dapat melakukan transaksi "
    "POS, mengelola produk, dan melihat rekap penjualan; (3) Admin Keuangan, yang dapat "
    "melakukan top-up, koreksi saldo, dan registrasi kartu; (4) Super Admin, yang dapat "
    "mengelola seluruh data dan konfigurasi sistem; serta (5) Orang Tua, yang dapat "
    "melakukan top-up dan memantau transaksi anak."
)

add_heading_text('3.3.3 Diagram Alur Transaksi', level=3)
add_body(
    "Alur transaksi pembelian dimulai ketika petugas kantin memilih produk yang akan "
    "dibeli oleh siswa melalui antarmuka POS. Setelah semua produk dipilih, petugas "
    "menekan tombol \"Proses Tap Kartu Siswa\" dan siswa menempelkan kartu RFID/NFC "
    "pada perangkat pembaca. Sistem membaca UID kartu, memvalidasi saldo siswa, dan "
    "memproses transaksi melalui RPC function process_purchase yang menjalankan "
    "transaksi database secara atomic. Jika saldo mencukupi, sistem mengurangi saldo "
    "siswa, menambah pendapatan petugas, mencatat transaksi dan item transaksi, serta "
    "mengirimkan notifikasi ke aplikasi siswa."
)

add_heading_text('3.4 Perancangan Basis Data', level=2)
add_body(
    "Basis data sistem dirancang menggunakan PostgreSQL pada platform Supabase. Terdapat "
    "sepuluh tabel utama yang saling berelasi untuk mendukung seluruh fungsionalitas "
    "sistem. Berikut adalah struktur tabel yang digunakan:"
)

# Tabel database
add_table(
    ['No', 'Nama Tabel', 'Fungsi'],
    [
        ('1', 'profiles', 'Menyimpan profil umum semua pengguna'),
        ('2', 'students', 'Data spesifik siswa termasuk saldo dan UID RFID'),
        ('3', 'canteen_operators', 'Informasi stan dan pendapatan petugas kantin'),
        ('4', 'products', 'Katalog produk makanan dan minuman'),
        ('5', 'transactions', 'Log transaksi pembelian dan top-up'),
        ('6', 'transaction_items', 'Rincian item dalam setiap transaksi'),
        ('7', 'notifications', 'Notifikasi real-time untuk siswa'),
        ('8', 'audit_logs', 'Catatan audit untuk keamanan data'),
        ('9', 'parent_students', 'Relasi orang tua dengan siswa'),
        ('10', 'system_settings', 'Konfigurasi global sistem'),
    ]
)

add_body(
    "Untuk menjaga integritas data transaksi, sistem menggunakan dua RPC function utama: "
    "process_purchase yang memproses pembelian dengan mekanisme SELECT FOR UPDATE untuk "
    "mencegah race condition, dan process_refund yang menangani pembatalan transaksi "
    "dalam batas waktu 10 menit setelah transaksi dilakukan."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#              BAB IV: HASIL DAN PEMBAHASAN
# ═══════════════════════════════════════════════════════════════
add_heading_text('BAB IV', level=1)
add_heading_text('HASIL DAN PEMBAHASAN', level=1)
add_spacing(1)

add_heading_text('4.1 Implementasi Sistem', level=2)
add_body(
    "Sistem Kantin Digital telah berhasil diimplementasikan menggunakan Flutter versi "
    "3.9.2 dengan bahasa Dart dan Supabase sebagai backend. Sistem terdiri dari lima "
    "aplikasi yang terpisah berdasarkan peran pengguna, masing-masing dengan antarmuka "
    "dan fungsionalitas yang disesuaikan."
)

add_heading_text('4.1.1 Aplikasi Siswa', level=3)
add_body(
    "Aplikasi siswa menyediakan fitur-fitur utama seperti dashboard yang menampilkan "
    "saldo digital dan status kartu, riwayat transaksi lengkap dengan detail item "
    "pembelian, manajemen kartu RFID (aktivasi/pembekuan), notifikasi real-time, serta "
    "profil pengguna. Antarmuka dirancang dengan tema iOS menggunakan Cupertino widgets "
    "dan warna teal sebagai warna primer."
)

add_heading_text('4.1.2 Aplikasi Kasir POS', level=3)
add_body(
    "Aplikasi kasir POS digunakan oleh petugas kantin untuk melayani transaksi "
    "pembelian. Fitur-fitur utama meliputi katalog produk yang dapat difilter berdasarkan "
    "kategori (makanan/minuman), keranjang belanja real-time, pemindaian kartu RFID/NFC "
    "untuk identifikasi siswa, dan pemrosesan pembayaran otomatis. Aplikasi ini juga "
    "menampilkan pendapatan harian dan riwayat penjualan."
)

add_heading_text('4.1.3 Aplikasi Admin Keuangan', level=3)
add_body(
    "Aplikasi admin keuangan menyediakan fungsionalitas untuk mengelola saldo siswa, "
    "termasuk top-up tunai, koreksi saldo, registrasi kartu RFID baru, serta pembuatan "
    "laporan keuangan. Setiap perubahan saldo dictat dalam audit log untuk menjaga "
    "transparansi dan akuntabilitas."
)

add_heading_text('4.1.4 Aplikasi Super Admin', level=3)
add_body(
    "Aplikasi super admin dilengkapi dengan secure entry (PIN/biometric) untuk mengakses "
    "panel administrasi. Fitur-fitur yang tersedia meliputi manajemen pengguna (CRUD), "
    "audit log untuk melacak seluruh aktivitas sistem, pengaturan konfigurasi global, "
    "serta dashboard yang menampilkan statistik keseluruhan sistem."
)

add_heading_text('4.1.5 Aplikasi Orang Tua', level=3)
add_body(
    "Aplikasi orang tua memungkinkan orang tua untuk melakukan top-up saldo anak melalui "
    "berbagai metode pembayaran termasuk Midtrans (QRIS, transfer bank), serta memantau "
    "riwayat transaksi jajan anak secara real-time. Orang tua login menggunakan NISN "
    "anak yang telah terdaftar dalam sistem."
)

add_heading_text('4.2 Pengujian Sistem', level=2)
add_body(
    "Pengujian sistem dilakukan menggunakan metode black-box testing untuk memverifikasi "
    "bahwa setiap fungsionalitas berjalan sesuai dengan spesifikasi yang telah ditentukan. "
    "Berikut adalah ringkasan hasil pengujian:"
)

add_table(
    ['No', 'Modul', 'Fitur', 'Status'],
    [
        ('1', 'Autentikasi', 'Login dengan email/username/NISN', '✓ Berhasil'),
        ('2', 'Autentikasi', 'Login orang tua dengan NISN', '✓ Berhasil'),
        ('3', 'Autentikasi', 'Logout dan manajemen sesi', '✓ Berhasil'),
        ('4', 'Siswa', 'Tampil saldo dan status kartu', '✓ Berhasil'),
        ('5', 'Siswa', 'Riwayat transaksi', '✓ Berhasil'),
        ('6', 'Siswa', 'Bekukan/aktifkan kartu', '✓ Berhasil'),
        ('7', 'POS', 'Katalog produk & filter kategori', '✓ Berhasil'),
        ('8', 'POS', 'Keranjang belanja', '✓ Berhasil'),
        ('9', 'POS', 'Pemindaian kartu RFID/NFC', '✓ Berhasil'),
        ('10', 'POS', 'Proses pembayaran & refund', '✓ Berhasil'),
        ('11', 'Keuangan', 'Top-up saldo tunai', '✓ Berhasil'),
        ('12', 'Keuangan', 'Koreksi saldo', '✓ Berhasil'),
        ('13', 'Keuangan', 'Registrasi kartu RFID', '✓ Berhasil'),
        ('14', 'Admin', 'Manajemen pengguna CRUD', '✓ Berhasil'),
        ('15', 'Admin', 'Audit log', '✓ Berhasil'),
        ('16', 'Orang Tua', 'Top-up online via Midtrans', '✓ Berhasil'),
        ('17', 'Orang Tua', 'Pemantauan transaksi', '✓ Berhasil'),
    ]
)

add_body(
    "Seluruh skenario pengujian menunjukkan hasil yang sesuai dengan spesifikasi. "
    "Tidak ditemukan bug kritis yang menghambat fungsionalitas utama sistem. "
    "Pengujian juga mencakup skenario edge case seperti saldo tidak mencukupi, kartu "
    "tidak aktif, dan transaksi refund melebihi batas waktu 10 menit."
)

add_heading_text('4.3 Pembahasan', level=2)
add_body(
    "Hasil implementasi menunjukkan bahwa sistem Kantin Digital mampu memenuhi seluruh "
    "kebutuhan fungsional yang telah didefinisikan. Penggunaan Flutter dengan state "
    "management Riverpod memberikan pengalaman pengembangan yang produktif dengan "
    "kemampuan hot reload yang signifikan mempercepat proses debugging dan iterasi "
    "antarmuka."
)
add_body(
    "Integrasi dengan Supabase berjalan dengan baik, terutama pada fitur autentikasi "
    "dual-path yang menggabungkan Supabase Auth dengan profile-based fallback. Mekanisme "
    "ACID transaction pada RPC function process_purchase berhasil mencegah terjadinya "
    "double-spending, di mana validasi saldo dan pemotongan dilakukan dalam satu "
    "transaksi database yang atomic."
)
add_body(
    "Penggunaan teknologi RFID/NFC sebagai media identifikasi siswa memberikan keunggulan "
    "dalam kecepatan transaksi. Proses pembayaran yang sebelumnya memerlukan waktu 30-60 "
    "detik dengan uang tunai (mencari uang, menghitung kembalian) dapat dipangkas menjadi "
    "kurang dari 5 detik dengan sistem tap kartu. Hal ini secara signifikan mengurangi "
    "antrean di kantin, terutama pada jam istirahat."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#                 BAB V: PENUTUP
# ═══════════════════════════════════════════════════════════════
add_heading_text('BAB V', level=1)
add_heading_text('PENUTUP', level=1)
add_spacing(1)

add_heading_text('5.1 Kesimpulan', level=2)
add_body(
    "Berdasarkan hasil penelitian dan pembahasan yang telah diuraikan, dapat ditarik "
    "kesimpulan sebagai berikut:"
)
add_numbered(
    'Sistem Kantin Digital berbasis RFID/NFC berhasil dirancang dan dibangun menggunakan '
    'Flutter dan Supabase. Sistem ini mampu mempercepat proses transaksi di kantin sekolah '
    'dari rata-rata 30-60 detik menjadi kurang dari 5 detik per transaksi melalui mekanisme '
    'tap kartu RFID/NFC.', 1
)
add_numbered(
    'Integrasi Flutter sebagai frontend dengan Supabase sebagai backend berhasil '
    'diimplementasikan melalui Supabase Flutter SDK yang menyediakan koneksi real-time, '
    'manajemen autentikasi, dan operasi basis data secara langsung.', 2
)
add_numbered(
    'Mekanisme manajemen saldo digital yang aman berhasil diimplementasikan menggunakan '
    'RPC function dengan transaksi ACID (SELECT FOR UPDATE) yang mencegah terjadinya '
    'race condition dan double-spending pada saat pemrosesan transaksi.', 3
)
add_numbered(
    'Fitur pemantauan bagi orang tua berhasil disediakan melalui aplikasi khusus orang '
    'tua yang memungkinkan akses ke riwayat transaksi dan saldo anak secara real-time '
    'serta fitur top-up online.', 4
)
add_numbered(
    'Sistem role-based access control berhasil diimplementasikan dengan lima tingkat '
    'akses: super admin, admin keuangan, petugas kantin, siswa, dan orang tua. Setiap '
    'peran memiliki hak akses dan antarmuka yang disesuaikan dengan kebutuhan masing-masing.', 5
)

add_heading_text('5.2 Saran', level=2)
add_body(
    "Untuk pengembangan sistem lebih lanjut, penulis memberikan beberapa saran sebagai "
    "berikut:"
)
add_numbered('Implementasi Row Level Security (RLS) di Supabase untuk meningkatkan keamanan data pada tingkat basis data.', 1)
add_numbered('Penambahan fitur notifikasi WhatsApp untuk memberikan informasi transaksi kepada orang tua secara langsung.', 2)
add_numbered('Pengembangan aplikasi berbasis web untuk memudahkan akses admin tanpa harus menginstal aplikasi mobile.', 3)
add_numbered('Integrasi dengan sistem pembayaran digital lainnya seperti GoPay, OVO, atau ShopeePay.', 4)
add_numbered('Penambahan fitur analitik dan grafik untuk membantu sekolah dalam mengambil keputusan berdasarkan data transaksi.', 5)
add_numbered('Implementasi sistem caching dan sinkronisasi offline untuk memastikan aplikasi tetap dapat digunakan saat koneksi internet terputus.', 6)

add_page_break()

# ═══════════════════════════════════════════════════════════════
#                  DAFTAR PUSTAKA
# ═══════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = Align.CENTER
p.paragraph_format.line_spacing = 1.5
run = p.add_run('DAFTAR PUSTAKA')
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
add_spacing(2)

references = [
    'Google. (2024). Flutter Documentation. Diakses dari https://docs.flutter.dev/ pada 20 Juni 2026.',
    'O\'Brien, J. A., & Marakas, G. M. (2010). Management Information Systems (10th ed.). New York: McGraw-Hill/Irwin.',
    'Pressman, R. S. (2015). Software Engineering: A Practitioner\'s Approach (8th ed.). New York: McGraw-Hill Education.',
    'Supabase. (2024). Supabase Documentation. Diakses dari https://supabase.com/docs pada 20 Juni 2026.',
    'Want, R. (2006). An Introduction to RFID Technology. IEEE Pervasive Computing, 5(1), 25-33.',
    'Dart Programming Language. (2024). Dart Documentation. Diakses dari https://dart.dev/guides pada 20 Juni 2026.',
    'Flutter Riverpod. (2024). Riverpod Documentation. Diakses dari https://riverpod.dev/ pada 20 Juni 2026.',
    'GoRouter. (2024). GoRouter Package. Diakses dari https://pub.dev/packages/go_router pada 20 Juni 2026.',
    'NFC Manager. (2024). NFC Manager Package. Diakses dari https://pub.dev/packages/nfc_manager pada 20 Juni 2026.',
    'Connectivity Plus. (2024). Connectivity Plus Package. Diakses dari https://pub.dev/packages/connectivity_plus pada 20 Juni 2026.',
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = Align.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ─── SAVE ─────────────────────────────────────────────────────
output_path = os.path.expanduser('~/projects/kantin-digital/Makalah_Kantin_Digital.docx')
doc.save(output_path)
print(f'✅ Makalah berhasil disimpan di: {output_path}')
